# Import
import os, re, json, tempfile
from urllib.parse import urlparse
from dotenv import load_dotenv  # Added for .env support

# Load environment variables
load_dotenv()

# Network and validation
import requests # to fetch web/pdf/caption files
import validators # to validate URL inputs

# UI Framework
import streamlit as st 

# langchin core pieces
from langchain_core.prompts import PromptTemplate
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_classic.chains.summarize import load_summarize_chain
from langchain_core.documents import Document

# loaders
from langchain_community.document_loaders import WebBaseLoader, YoutubeLoader, PyPDFLoader

# LLM
from langchain_groq import ChatGroq

# Youtube caption edge case and fallbacks
from youtube_transcript_api._errors import TranscriptsDisabled, NoTranscriptFound, VideoUnavailable
from yt_dlp import YoutubeDL

# New imports for additional features
from PIL import Image
import pytesseract
import docx
from io import StringIO

# Configuration
class Config:
    """Configuration class for managing app settings"""
    DEFAULT_CHUNK_SIZE = 1600
    DEFAULT_CHUNK_OVERLAP = 150
    DEFAULT_TEMPERATURE = 0.2
    DEFAULT_OUTPUT_LENGTH = 300
    SUPPORTED_LANGUAGES = ["English", "Urdu", "Roman Urdu", "Auto"]
    SUPPORTED_MODELS = ["llama-3.1-8b-instant", "openai/gpt-oss-20b", "openai/gpt-oss-120b"]  # Removed gemma2-9b-it
    SUPPORTED_CONTENT_TYPES = ["Website", "YouTube", "PDF", "Text File", "DOCX", "Image (OCR)"]
    SUMMARY_TEMPLATES = {
        "Default": "Standard summary",
        "Academic": "Focus on methodology, findings, and conclusions",
        "Business": "Focus on key metrics, recommendations, and action items",
        "Technical": "Focus on specifications, features, and technical details",
        "News": "Focus on key events, facts, and implications"
    }

# Minimal Page setup
st.set_page_config(
    page_title="OmniBrief - AI Summarizer", 
    page_icon="🧠", 
    layout="wide",
    initial_sidebar_state="expanded"
)

st.title("🧠 OmniBrief - Summarizer")
st.caption("Build with Streamlit + Langchain + Groq")

# Sidebar Configuration
with st.sidebar:
    st.header("🔑 API & Model Configuration")
    
    # API Key with better environment variable handling
    default_api_key = os.getenv("GROQ_API_KEY", "")
    groq_api_key = st.text_input(
        "Groq API Key", 
        type="password", 
        value=default_api_key,
        help="Enter your Groq API key or set GROQ_API_KEY in .env file"
    )
    
    # Model selection with better organization
    col1, col2 = st.columns([3, 1])
    with col1:
        model = st.selectbox(
            "Groq Model",
            Config.SUPPORTED_MODELS,
            index=0,
            help="Select a Groq model from the available options"
        )
    with col2:
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("🔄", help="Refresh models"):
            st.rerun()
    
    custom_model = st.text_input(
        "Custom Model (optional)", 
        help="Override the selected model with a custom model name"
    )

    st.header("🧠 Generation Settings")
    temperature = st.slider(
        "Temperature", 
        0.0, 1.0, Config.DEFAULT_TEMPERATURE, 0.05,
        help="Controls creativity: lower = more deterministic, higher = more creative"
    )
    out_len = st.slider(
        "Target Summary Length (words)", 
        50, 1000, Config.DEFAULT_OUTPUT_LENGTH, 10,
        help="Approximate word count for the final summary"
    )

    st.header("✍🏻 Output Style")
    out_style = st.selectbox("Output Format", ["Bullets", "Paragraph", "Both"])
    tone = st.selectbox("Tone", ["Neutral", "Formal", "Casual", "Executive Brief"])
    out_lang = st.selectbox("Language", Config.SUPPORTED_LANGUAGES)
    
    # New: Summary Template
    summary_template = st.selectbox(
        "Summary Template",
        list(Config.SUMMARY_TEMPLATES.keys()),
        help="Choose a template for specific summary types"
    )

    st.header("⚙️ Processing Settings")
    chain_mode = st.radio(
        "Processing Method", 
        ["Auto", "Stuff", "Map-Reduce"],
        index=0,
        help="Auto: Choose automatically, Stuff: Single pass, Map-Reduce: For long documents"
    )
    
    chunk_size = st.slider(
        "Chunk Size (characters)", 
        500, 5000, Config.DEFAULT_CHUNK_SIZE, 50,
        help="Size of text chunks for processing"
    )
    chunk_overlap = st.slider(
        "Chunk Overlap (characters)", 
        0, 1000, Config.DEFAULT_CHUNK_OVERLAP, 10,
        help="Overlap between chunks to maintain context"
    )
    max_map_chunks = st.slider(
        "Max Chunks for Processing", 
        5, 100, 28, 1,
        help="Maximum number of chunks to process (prevents timeout)"
    )

    st.header("🎯 Additional Features")
    show_preview = st.checkbox("Show Source Preview", value=True)
    want_outline = st.checkbox("Include Outline", value=True)
    want_keywords = st.checkbox("Extract Keywords & Hashtags", value=True)
    enable_translation = st.checkbox("Enable Translation", value=False)
    save_template = st.checkbox("Save Summary as Template", value=False)
    
    # Custom Instructions
    custom_instructions = st.text_area(
        "Custom Instructions",
        placeholder="Add any specific instructions for the summary...",
        help="Custom guidelines for how the summary should be generated"
    )

# Main Input Section
st.header("📥 Input Source")

# Input Method Selection
input_method = st.radio(
    "Input Method",
    ["Single URL", "Multiple URLs", "File Upload", "Direct Text"],
    horizontal=True
)

documents = []
metadata = {"type": None, "source": None, "title": None}

if input_method == "Single URL":
    url = st.text_input(
        "Paste URL", 
        placeholder="Enter website, YouTube, or PDF URL...",
        help="Supports: Websites, YouTube videos, PDF URLs"
    )
    
elif input_method == "Multiple URLs":
    urls = st.text_area(
        "Paste Multiple URLs (one per line)",
        placeholder="Enter multiple URLs...\nhttps://example1.com\nhttps://example2.com",
        help="One URL per line",
        height=100
    )
    
elif input_method == "File Upload":
    uploaded = st.file_uploader(
        "Upload File", 
        type=["pdf", "txt", "md", "docx", "png", "jpg", "jpeg"],
        help="Upload PDF, Text, DOCX, or Image files"
    )
    
elif input_method == "Direct Text":
    direct_text = st.text_area(
        "Paste Text Directly", 
        placeholder="Paste your text here...",
        height=200,
        help="Direct text input for summarization"
    )

# Utility Functions
class ContentProcessor:
    """Handles content processing and validation"""
    
    @staticmethod
    def is_youtube(url: str) -> bool:
        """Check if URL is from YouTube"""
        try:
            netloc = urlparse(url).netloc.lower()
            return any(host in netloc for host in ["youtube.com", "youtu.be"])
        except Exception:
            return False
    
    @staticmethod
    def get_content_type(url: str, timeout: int = 10) -> str | None:
        """Get content type of URL"""
        try:
            response = requests.head(
                url, 
                allow_redirects=True, 
                timeout=timeout, 
                headers={"User-Agent": "Mozilla/5.0"}
            )
            return (response.headers.get("Content-Type") or "").lower()
        except Exception:
            return None
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Clean and normalize text"""
        # Remove audio cues and extra spaces
        text = re.sub(r"\[(?:music|applause|laughter|.*?)\]", " ", text, flags=re.I)
        text = re.sub(r"\s+", " ", text)
        return text.strip()
    
    @staticmethod
    def validate_inputs(groq_api_key: str, input_method: str, url: str, urls: str, uploaded_file, direct_text: str) -> bool:
        """Validate user inputs"""
        if not groq_api_key.strip():
            st.error("❌ Please provide your Groq API Key in the sidebar")
            return False
        
        if input_method == "Single URL" and not url.strip():
            st.error("❌ Please provide a URL")
            return False
        
        if input_method == "Multiple URLs" and not urls.strip():
            st.error("❌ Please provide at least one URL")
            return False
            
        if input_method == "File Upload" and uploaded_file is None:
            st.error("❌ Please upload a file")
            return False
            
        if input_method == "Direct Text" and not direct_text.strip():
            st.error("❌ Please enter some text")
            return False
        
        if input_method == "Single URL" and url.strip() and not validators.url(url):
            st.error("❌ Please enter a valid URL")
            return False
            
        if input_method == "Multiple URLs" and urls.strip():
            url_list = [u.strip() for u in urls.split('\n') if u.strip()]
            for u in url_list:
                if not validators.url(u):
                    st.error(f"❌ Invalid URL: {u}")
                    return False
                    
        return True

class YouTubeProcessor:
    """Specialized processor for YouTube content"""
    
    @staticmethod
    def process_captions(caption_data: str, content_type: str) -> str:
        """Process YouTube captions from different formats"""
        if "application/json" in content_type or caption_data.strip().startswith("{"):
            return YouTubeProcessor._json_to_text(caption_data)
        elif "text/vtt" in content_type:
            return YouTubeProcessor._vtt_to_text(caption_data)
        else:
            return ContentProcessor.clean_text(caption_data)
    
    @staticmethod
    def _json_to_text(json_data: str) -> str:
        """Convert JSON caption format to text"""
        try:
            data = json.loads(json_data)
            lines = []
            for event in data.get("events", []):
                for segment in event.get("segs", []) or []:
                    text = segment.get("utf8", "")
                    if text:
                        lines.append(text.replace("\n", " "))
            return ContentProcessor.clean_text(" ".join(lines))
        except Exception:
            return ContentProcessor.clean_text(json_data)
    
    @staticmethod
    def _vtt_to_text(vtt_data: str) -> str:
        """Convert VTT caption format to text"""
        lines = []
        for line in vtt_data.splitlines():
            stripped = line.strip()
            # Skip timestamps and headers
            if ("-->" in stripped) or stripped.isdigit() or stripped.upper().startswith("WEBVTT"):
                continue
            if stripped:
                lines.append(stripped)
        return ContentProcessor.clean_text(" ".join(lines))

class FileProcessor:
    """Handles various file types"""
    
    @staticmethod
    def process_text_file(uploaded_file) -> str:
        """Process text files"""
        try:
            stringio = StringIO(uploaded_file.getvalue().decode("utf-8"))
            text = stringio.read()
            return ContentProcessor.clean_text(text)
        except Exception as e:
            raise RuntimeError(f"Failed to read text file: {e}")
    
    @staticmethod
    def process_docx_file(uploaded_file) -> str:
        """Process DOCX files"""
        try:
            doc = docx.Document(uploaded_file)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return ContentProcessor.clean_text(text)
        except Exception as e:
            raise RuntimeError(f"Failed to read DOCX file: {e}")
    
    @staticmethod
    def process_image_file(uploaded_file) -> str:
        """Process image files with OCR"""
        try:
            image = Image.open(uploaded_file)
            text = pytesseract.image_to_string(image)
            return ContentProcessor.clean_text(text)
        except Exception as e:
            raise RuntimeError(f"Failed to process image with OCR: {e}")

class LLMProcessor:
    """Handles LLM-related operations"""
    
    @staticmethod
    def build_llm(groq_api_key: str, model: str, temperature: float) -> ChatGroq:
        """Initialize Groq LLM"""
        chosen_model = custom_model.strip() if custom_model else model
        return ChatGroq(
            model=chosen_model, 
            groq_api_key=groq_api_key, 
            temperature=temperature
        )
    
    @staticmethod
    def build_prompts(out_len: int, out_style: str, tone: str, 
                     want_outline: bool, want_keywords: bool, out_lang: str,
                     summary_template: str, custom_instructions: str):
        """Build map and combine prompts for summarization"""
        
        # Template-specific instructions
        template_instructions = Config.SUMMARY_TEMPLATES.get(summary_template, "")
        
        # Map prompt for initial chunk processing
        map_template = f"""
        Analyze the following text and extract the key points. Focus on:
        - Main ideas and arguments
        - Important facts and data
        - Key conclusions or recommendations
        {f"- {template_instructions}" if template_instructions and summary_template != "Default" else ""}
        {f"- Additional instructions: {custom_instructions}" if custom_instructions.strip() else ""}
        
        Be concise and objective.
        
        TEXT:
        {{text}}
        
        KEY POINTS:
        """
        map_prompt = PromptTemplate(template=map_template, input_variables=["text"])
        
        # Style and tone mappings
        style_instructions = {
            "Bullets": "Present the summary as bullet points only",
            "Paragraph": "Present the summary as a cohesive paragraph only",
            "Both": "Start with concise bullet points followed by a cohesive paragraph",
        }
        
        tone_instructions = {
            "Neutral": "neutral and objective",
            "Formal": "formal and professional",
            "Casual": "casual and conversational", 
            "Executive Brief": "executive-level, action-oriented, and concise",
        }
        
        # Language instruction
        lang_instruction = "Use the same language as the source content." if out_lang == "Auto" else f"Write in {out_lang}."
        
        # Additional sections
        extra_sections = []
        if want_outline:
            extra_sections.append("Provide a structured outline with main sections.")
        if want_keywords:
            extra_sections.append("Extract 8-12 relevant keywords and 5-8 hashtags.")
        
        extra_instructions = "\n".join(f"- {section}" for section in extra_sections)
        
        # Combine prompt for final summary
        combine_template = f"""
        You are creating a comprehensive summary from multiple text segments.
        
        REQUIREMENTS:
        - Target length: {out_len} words
        - Style: {style_instructions[out_style]}
        - Tone: {tone_instructions[tone]}
        - Language: {lang_instruction}
        {f"- Template focus: {template_instructions}" if template_instructions and summary_template != "Default" else ""}
        {f"- Custom instructions: {custom_instructions}" if custom_instructions.strip() else ""}
        - Be accurate and faithful to the source material
        - Clearly distinguish between facts and opinions
        - Avoid repetition and focus on essential information
        
        ADDITIONAL SECTIONS:
        {extra_instructions if extra_sections else "None"}
        
        INPUT SUMMARIES:
        {{text}}
        
        FINAL SUMMARY:
        """
        combine_prompt = PromptTemplate(template=combine_template, input_variables=["text"])
        
        return map_prompt, combine_prompt

class DocumentProcessor:
    """Handles document loading and processing"""
    
    @staticmethod
    def load_youtube_content(url: str):
        """Load content from YouTube URL with enhanced error handling"""
        # Primary method: youtube_transcript_api with better configuration
        try:
            loader = YoutubeLoader.from_youtube_url(
                url,
                add_video_info=False,  # Simplified to avoid issues
                language=["en", "en-US", "en-GB", "ur", "hi"],
                translation="en",  # Add translation fallback
            )
            docs = loader.load()
            if docs and any((doc.page_content or "").strip() for doc in docs):
                return docs, {"type": "youtube", "source": url}
        except Exception as e:
            st.warning(f"Primary YouTube loading failed: {e}. Trying fallback...")
        
        # Enhanced fallback: yt-dlp for captions
        try:
            return DocumentProcessor._youtube_enhanced_fallback(url)
        except Exception as e:
            # Final fallback: extract metadata only
            try:
                return DocumentProcessor._youtube_metadata_fallback(url)
            except Exception as final_e:
                raise RuntimeError(f"All YouTube loading methods failed: {final_e}")
    
    @staticmethod
    def _youtube_enhanced_fallback(url: str):
        """Enhanced fallback method for YouTube content loading"""
        try:
            ydl_opts = {
                "skip_download": True,
                "quiet": True,
                "writesubtitles": True,
                "writeautomaticsub": True,
                "subtitleslangs": ["en", "en-US", "en-GB", "ur", "hi"],
                "subtitlesformat": "json3"  # Prefer JSON format
            }
            
            with YoutubeDL(ydl_opts) as ydl:
                info = ydl.extract_info(url, download=False)
                
                # Show available subtitle languages
                available_subs = list(info.get('subtitles', {}).keys()) + list(info.get('automatic_captions', {}).keys())
                if available_subs:
                    st.info(f"🎯 Available subtitle languages: {available_subs}")
                
                # Try different subtitle formats and languages
                subtitle_formats = ["json3", "vtt", "ttml", "srv3", "srv2", "srv1"]
                languages = ["en", "en-US", "en-GB", "a.en", "ur", "hi"]
                
                for lang in languages:
                    for fmt in subtitle_formats:
                        try:
                            caption_url = None
                            
                            # Try manual subtitles first
                            if info.get('subtitles') and lang in info['subtitles']:
                                for sub in info['subtitles'][lang]:
                                    if sub.get('ext') == fmt:
                                        caption_url = sub['url']
                                        break
                            
                            # Try automatic captions
                            if not caption_url and info.get('automatic_captions') and lang in info['automatic_captions']:
                                for sub in info['automatic_captions'][lang]:
                                    if sub.get('ext') == fmt:
                                        caption_url = sub['url']
                                        break
                            
                            if caption_url:
                                # Fetch captions
                                response = requests.get(caption_url, timeout=30, headers={
                                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
                                })
                                response.raise_for_status()
                                
                                caption_text = YouTubeProcessor.process_captions(response.text, f"application/{fmt}")
                                if caption_text and len(caption_text.strip()) > 100:  # Minimum content check
                                    st.success(f"✅ Using {lang} subtitles in {fmt} format")
                                    return [Document(page_content=caption_text, metadata={"source": url})], {"type": "youtube_fallback"}
                                    
                        except Exception:
                            continue
                
                raise RuntimeError("No usable captions found")
                
        except Exception as e:
            raise RuntimeError(f"YouTube enhanced fallback failed: {e}")
    
    @staticmethod
    def _youtube_metadata_fallback(url: str):
        """Final fallback using YouTube video metadata"""
        try:
            ydl_opts = {
                "skip_download": True,
                "quiet": True,
                "no_warnings": True,
            }
            
            with YoutubeDL(ydl_opts) as ydl:
                info = ydl.extract_info(url, download=False)
                
                video_title = info.get('title', 'Unknown Title')
                video_description = info.get('description', '')
                video_duration = info.get('duration', 0)
                video_uploader = info.get('uploader', 'Unknown Uploader')
                
                # Create basic content from metadata
                content_parts = [
                    f"Video Title: {video_title}",
                    f"Uploader: {video_uploader}",
                    f"Duration: {video_duration} seconds",
                ]
                
                if video_description:
                    content_parts.append(f"\nDescription:\n{video_description}")
                
                full_content = "\n".join(content_parts)
                
                if len(full_content.strip()) > 50:
                    st.warning("⚠️ Using video metadata only (no captions available)")
                    return [Document(page_content=full_content, metadata={"source": url})], {"type": "youtube_metadata"}
                else:
                    raise RuntimeError("Insufficient metadata available")
                    
        except Exception as e:
            raise RuntimeError(f"YouTube metadata fallback failed: {e}")
    
    @staticmethod
    def load_web_content(url: str, chunk_size: int, chunk_overlap: int):
        """Load content from web URL"""
        metadata = {"source": url, "type": "webpage", "title": None}
        
        try:
            loader = WebBaseLoader([url])
            docs = loader.load()
            if docs and docs[0].metadata.get("title"):
                metadata["title"] = docs[0].metadata["title"]
        except Exception as e:
            st.warning(f"WebBaseLoader failed: {e}. Using fallback method...")
            # Fallback: simple requests + regex
            try:
                response = requests.get(url, timeout=20, headers={"User-Agent": "Mozilla/5.0"})
                text = re.sub(r"<[^>]+>", " ", response.text)
                docs = [Document(page_content=text, metadata={"source": url})]
            except Exception as e:
                raise RuntimeError(f"Failed to load web content: {e}")
        
        return DocumentProcessor._chunk_documents(docs, chunk_size, chunk_overlap), metadata
    
    @staticmethod
    def load_pdf_content(url: str, chunk_size: int, chunk_overlap: int):
        """Load content from PDF URL"""
        metadata = {"source": url, "type": "pdf", "title": None}
        
        try:
            with requests.get(url, stream=True, timeout=20, headers={"User-Agent": "Mozilla/5.0"}) as response:
                response.raise_for_status()
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
                    for chunk in response.iter_content(chunk_size=8192):
                        if chunk:
                            temp_file.write(chunk)
                    temp_path = temp_file.name
            
            loader = PyPDFLoader(temp_path)
            docs = loader.load()
            
            # Clean up temporary file
            try:
                os.unlink(temp_path)
            except Exception:
                pass
                
            return DocumentProcessor._chunk_documents(docs, chunk_size, chunk_overlap), metadata
            
        except Exception as e:
            raise RuntimeError(f"Failed to load PDF from URL: {e}")
    
    @staticmethod
    def load_uploaded_pdf(uploaded_file, chunk_size: int, chunk_overlap: int):
        """Load content from uploaded PDF"""
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
            temp_file.write(uploaded_file.getbuffer())
            temp_path = temp_file.name
        
        try:
            loader = PyPDFLoader(temp_path)
            docs = loader.load()
            return DocumentProcessor._chunk_documents(docs, chunk_size, chunk_overlap)
        finally:
            # Clean up temporary file
            try:
                os.unlink(temp_path)
            except Exception:
                pass
    
    @staticmethod
    def load_multiple_urls(urls: str, chunk_size: int, chunk_overlap: int):
        """Load content from multiple URLs"""
        url_list = [u.strip() for u in urls.split('\n') if u.strip()]
        all_docs = []
        metadata = {"type": "multiple_urls", "sources": url_list}
        
        for url in url_list:
            try:
                if ContentProcessor.is_youtube(url):
                    docs, _ = DocumentProcessor.load_youtube_content(url)
                elif "pdf" in ContentProcessor.get_content_type(url) or url.lower().endswith(".pdf"):
                    docs, _ = DocumentProcessor.load_pdf_content(url, chunk_size, chunk_overlap)
                else:
                    docs, _ = DocumentProcessor.load_web_content(url, chunk_size, chunk_overlap)
                
                all_docs.extend(docs)
                st.success(f"✅ Loaded: {url}")
                
            except Exception as e:
                st.warning(f"⚠️ Failed to load {url}: {e}")
                continue
        
        if not all_docs:
            raise RuntimeError("No content could be loaded from any URL")
            
        return DocumentProcessor._chunk_documents(all_docs, chunk_size, chunk_overlap), metadata
    
    @staticmethod
    def process_uploaded_file(uploaded_file, chunk_size: int, chunk_overlap: int):
        """Process various uploaded file types"""
        file_extension = uploaded_file.name.lower().split('.')[-1]
        
        if file_extension == 'pdf':
            return DocumentProcessor.load_uploaded_pdf(uploaded_file, chunk_size, chunk_overlap)
        
        elif file_extension in ['txt', 'md']:
            text = FileProcessor.process_text_file(uploaded_file)
            return [Document(page_content=text, metadata={"source": uploaded_file.name})]
        
        elif file_extension == 'docx':
            text = FileProcessor.process_docx_file(uploaded_file)
            return [Document(page_content=text, metadata={"source": uploaded_file.name})]
        
        elif file_extension in ['png', 'jpg', 'jpeg']:
            text = FileProcessor.process_image_file(uploaded_file)
            return [Document(page_content=text, metadata={"source": uploaded_file.name})]
        
        else:
            raise RuntimeError(f"Unsupported file type: {file_extension}")
    
    @staticmethod
    def process_direct_text(direct_text: str):
        """Process direct text input"""
        cleaned_text = ContentProcessor.clean_text(direct_text)
        return [Document(page_content=cleaned_text, metadata={"source": "direct_text"})]
    
    @staticmethod
    def _chunk_documents(docs, chunk_size: int, chunk_overlap: int):
        """Chunk documents if they exceed size threshold"""
        total_size = sum(len(doc.page_content or "") for doc in docs)
        if total_size > chunk_size * 1.5:
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                separators=["\n\n", "\n", ".", "?", "!", " "],
            )
            chunked_docs = []
            for doc in docs:
                chunked_docs.extend(splitter.split_documents([doc]))
            return chunked_docs
        return docs

class SummarizationEngine:
    """Main summarization engine"""
    
    @staticmethod
    def choose_chain_type(mode: str, documents: list) -> str:
        """Choose the appropriate chain type"""
        if mode != "Auto":
            return mode.lower().replace("-", "_")
        
        total_chars = sum(len(doc.page_content or "") for doc in documents)
        return "map_reduce" if total_chars > 15000 else "stuff"
    
    @staticmethod
    def sample_documents(documents: list, max_chunks: int) -> list:
        """Evenly sample documents to avoid exceeding context limits"""
        if max_chunks >= len(documents):
            return documents
        
        indices = [round(i * (len(documents) - 1) / (max_chunks - 1)) for i in range(max_chunks)]
        return [documents[i] for i in indices]
    
    @staticmethod
    def run_summarization(llm, documents, map_prompt: PromptTemplate, 
                         combine_prompt: PromptTemplate, mode: str, max_chunks: int) -> str:
        """Execute the summarization chain"""
        mode = mode.lower().replace("-", "_")
        
        # Sample documents if needed for map_reduce
        if mode == "map_reduce" and len(documents) > max_chunks:
            documents = SummarizationEngine.sample_documents(documents, max_chunks)
            st.info(f"📊 Long document: sampled {max_chunks} chunks for processing.")
        
        try:
            # Load appropriate chain with updated interface
            if mode == "stuff":
                chain = load_summarize_chain(
                    llm, 
                    chain_type="stuff", 
                    prompt=combine_prompt,
                    verbose=False
                )
            else:
                chain = load_summarize_chain(
                    llm,
                    chain_type="map_reduce",
                    map_prompt=map_prompt,
                    combine_prompt=combine_prompt,
                    verbose=False
                )
            
            # Use the correct invocation method for the chain
            try:
                # Try the modern interface first
                result = chain.invoke({"input_documents": documents})
                return result.get("output_text", str(result))
            except Exception as e:
                # Fallback to older interface
                st.warning(f"Modern interface failed, trying fallback: {e}")
                try:
                    return chain.run(input_documents=documents)
                except Exception as fallback_e:
                    # Final fallback - manual processing
                    st.warning(f"Chain execution failed, using manual processing: {fallback_e}")
                    return SummarizationEngine._manual_summarization(llm, documents, combine_prompt)
                    
        except Exception as e:
            st.error(f"Chain execution failed: {e}")
            # Final fallback to manual processing
            return SummarizationEngine._manual_summarization(llm, documents, combine_prompt)
    
    @staticmethod
    def _manual_summarization(llm, documents, combine_prompt: PromptTemplate) -> str:
        """Manual summarization as fallback when chain fails"""
        try:
            # Combine all document content
            combined_text = "\n\n".join([doc.page_content for doc in documents if doc.page_content])
            
            # Truncate if too long (keep it reasonable)
            if len(combined_text) > 100000:
                combined_text = combined_text[:100000] + "... [content truncated]"
            
            # Use the combine prompt directly
            prompt_text = combine_prompt.format(text=combined_text)
            
            # Get response from LLM
            response = llm.invoke(prompt_text)
            
            return response.content if hasattr(response, 'content') else str(response)
            
        except Exception as e:
            raise RuntimeError(f"Manual summarization also failed: {e}")

# Main execution flow
def main():
    """Main application flow"""
    st.markdown("---")
    st.markdown("### 🚀 Generate Summary")
    
    if st.button("Summarize Content", type="primary", use_container_width=True):
        # Validate inputs based on input method
        if input_method == "Single URL":
            valid = ContentProcessor.validate_inputs(groq_api_key, input_method, url, "", None, "")
        elif input_method == "Multiple URLs":
            valid = ContentProcessor.validate_inputs(groq_api_key, input_method, "", urls, None, "")
        elif input_method == "File Upload":
            valid = ContentProcessor.validate_inputs(groq_api_key, input_method, "", "", uploaded, "")
        else:  # Direct Text
            valid = ContentProcessor.validate_inputs(groq_api_key, input_method, "", "", None, direct_text)
        
        if not valid:
            return
        
        documents = []
        metadata = {"type": None, "source": None, "title": None}
        
        try:
            # Load content based on input type
            with st.status("Loading content...", expanded=True) as status:
                status.update(label="🔄 Processing input source...")
                
                if input_method == "File Upload" and uploaded is not None:
                    documents = DocumentProcessor.process_uploaded_file(uploaded, chunk_size, chunk_overlap)
                    metadata.update({"type": uploaded.name.split('.')[-1], "source": uploaded.name})
                    status.update(label=f"✅ {uploaded.name} processed successfully")
                    
                elif input_method == "Single URL" and url.strip():
                    if ContentProcessor.is_youtube(url):
                        documents, yt_meta = DocumentProcessor.load_youtube_content(url)
                        metadata.update(yt_meta)
                        status.update(label="✅ YouTube content loaded successfully")
                    elif "pdf" in ContentProcessor.get_content_type(url) or url.lower().endswith(".pdf"):
                        documents, pdf_meta = DocumentProcessor.load_pdf_content(url, chunk_size, chunk_overlap)
                        metadata.update(pdf_meta)
                        status.update(label="✅ PDF content loaded successfully")
                    else:
                        documents, web_meta = DocumentProcessor.load_web_content(url, chunk_size, chunk_overlap)
                        metadata.update(web_meta)
                        status.update(label="✅ Web content loaded successfully")
                
                elif input_method == "Multiple URLs" and urls.strip():
                    documents, multi_meta = DocumentProcessor.load_multiple_urls(urls, chunk_size, chunk_overlap)
                    metadata.update(multi_meta)
                    status.update(label="✅ Multiple URLs processed successfully")
                
                elif input_method == "Direct Text" and direct_text.strip():
                    documents = DocumentProcessor.process_direct_text(direct_text)
                    metadata.update({"type": "direct_text", "source": "User Input"})
                    status.update(label="✅ Direct text processed successfully")
                
                # Validate extracted content
                if not documents or not any((doc.page_content or "").strip() for doc in documents):
                    st.error("❌ No readable text could be extracted from the source.")
                    return
            
            # Show source preview
            if show_preview and documents:
                with st.expander("🔍 Source Preview", expanded=False):
                    st.write(f"**Content Type:** `{metadata.get('type', 'Unknown')}`")
                    if metadata.get("title"):
                        st.write(f"**Title:** {metadata['title']}")
                    if metadata.get("source"):
                        st.write(f"**Source:** {metadata['source']}")
                    if metadata.get("sources"):
                        st.write(f"**Sources:** {len(metadata['sources'])} URLs")
                    
                    preview_text = "".join(doc.page_content or "" for doc in documents[:3])[:1500].strip()
                    st.text_area(
                        "First ~1500 characters", 
                        preview_text, 
                        height=200,
                        help="Preview of the extracted content"
                    )
            
            # Initialize LLM and prompts
            with st.status("Initializing summarization...", expanded=True) as status:
                status.update(label="🤖 Initializing AI model...")
                llm = LLMProcessor.build_llm(groq_api_key, model, temperature)
                
                status.update(label="📝 Preparing prompts...")
                map_prompt, combine_prompt = LLMProcessor.build_prompts(
                    out_len, out_style, tone, want_outline, want_keywords, out_lang,
                    summary_template, custom_instructions
                )
                
                status.update(label="⚙️ Configuring processing chain...")
                chain_type = SummarizationEngine.choose_chain_type(chain_mode, documents)
            
            # Execute summarization
            with st.status("Generating summary...", expanded=True) as status:
                status.update(label=f"🧠 Processing with {custom_model or model} ({chain_type})...")
                summary = SummarizationEngine.run_summarization(
                    llm, documents, map_prompt, combine_prompt, chain_type, max_map_chunks
                )
                status.update(label="✅ Summary generated successfully!")
            
            # Display results
            st.success("Summary Completed!")
            st.subheader("📋 Summary")
            
            # Summary display with better formatting
            st.write(summary)
            
            # Export options
            export_col1, export_col2, export_col3 = st.columns(3)
            with export_col1:
                st.download_button(
                    "📥 Download as Text",
                    data=summary,
                    file_name="summary.txt",
                    mime="text/plain",
                    use_container_width=True
                )
            with export_col2:
                st.download_button(
                    "📥 Download as Markdown", 
                    data=f"# Summary\n\n{summary}\n",
                    file_name="summary.md",
                    mime="text/markdown", 
                    use_container_width=True
                )
            with export_col3:
                st.download_button(
                    "📥 Download as JSON",
                    data=json.dumps({
                        "summary": summary,
                        "metadata": metadata,
                        "settings": {
                            "model": custom_model or model,
                            "template": summary_template,
                            "length": out_len
                        }
                    }, indent=2),
                    file_name="summary.json",
                    mime="application/json",
                    use_container_width=True
                )
                
        except Exception as e:
            st.error(f"❌ Processing failed: {str(e)}")
            st.info("💡 Check the troubleshooting notes below for possible solutions.")

# Run the application
if __name__ == "__main__":
    main()

# Troubleshooting and information
with st.expander("🚨 Troubleshooting & Information", expanded=False):
    st.markdown("""
    ### ✅ Supported Sources
    - **Websites**: Most public websites with readable text content
    - **YouTube**: Videos with captions (manual or automatic)  
    - **PDFs**: Direct links and uploaded files with selectable text
    - **Text Files**: TXT, MD files
    - **DOCX**: Microsoft Word documents
    - **Images**: PNG, JPG with OCR text extraction
    - **Direct Text**: Copy-paste text directly
    - **Multiple URLs**: Process multiple sources at once

    ### ⚠️ Common Issues
    - **Login Walls**: Cannot access password-protected content
    - **JavaScript Heavy Sites**: Some modern websites may not load properly
    - **Scanned PDFs**: Image-based PDFs require OCR (not supported)
    - **Content Blocks**: Some sites block automated access
    - **Long Videos**: Very long videos may exceed processing limits

    ### 🛠️ Performance Tips
    - **For long documents**: Use Map-Reduce method
    - **Adjust chunk size**: Smaller chunks for better accuracy, larger for speed
    - **Limit max chunks**: Prevents timeout on very long content
    - **Lower temperature**: For more factual, consistent results

    ### 🔧 Technical Notes
    - Uses LangChain for document processing
    - Groq API for fast LLM inference
    - Automatic fallback methods for robust loading
    - Configurable processing parameters

    ### 🎯 YouTube Specific Notes
    - **Captions Required**: Works best with videos that have English captions
    - **Fallback Methods**: Uses multiple approaches to extract content
    - **Metadata Fallback**: If no captions, uses video title and description
    - **Language Support**: Supports multiple subtitle languages

    ### 📁 File Support
    - **PDF**: Text-based PDFs (not scanned images)
    - **TXT/MD**: Plain text files
    - **DOCX**: Microsoft Word documents  
    - **Images**: OCR text extraction from PNG/JPG
    """)